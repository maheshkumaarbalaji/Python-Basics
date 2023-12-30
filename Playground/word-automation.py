"""
    1. Check if the folder word exists, if it is not, create a new one
    2. create a new document object add paragraphs, styling and tables
    3. Save the Word file inside the folder created with the current timestamp as name.
"""

import docx
from datetime import datetime
import os

Current_DateTime = str(datetime.now())
Current_DateTime = Current_DateTime.replace(" ", "-")
Current_DateTime = Current_DateTime.replace(".", "-")
Current_DateTime = Current_DateTime.replace(":", "-")
New_File_Name = "Word-File-" + Current_DateTime
Current_Folder_Contents = os.listdir(os.getcwd())
print(f"The current working directory is {os.getcwd()} and its contents are:", Current_Folder_Contents)
if "Word Documents" not in Current_Folder_Contents:
    os.mkdir("Word Documents")
    print("New folder 'Word Documents' has been created in the current working directory.")
else:
    print("The folder 'Word Documents' already exists in the current working directory.")
document = docx.Document()
document.add_heading("A Sample Word Document", 1)
document.add_paragraph(" ")
document.add_paragraph("This is a test paragraph to test the creation of a word document through python code")
document.add_picture("./Images/my-image-one.jpeg")
document.save("./Word Documents/" + New_File_Name + ".docx")
print(f"A new word document {New_File_Name + ".docx"} has been created in the Word Documents folder.")
